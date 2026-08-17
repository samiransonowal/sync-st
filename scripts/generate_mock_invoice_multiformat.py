#!/usr/bin/env python3
"""
ST-IN-gen — Multi-Format Mock Invoice Generator & Google Drive Ingestion (WIP)
Script: scripts/generate_mock_invoice_multiformat.py

Generates a complete mock invoice across 4 distinct formats:
1. Vector HTML (Web Invoice)
2. Vector PDF (Print Ready A4)
3. Excel Spreadsheet (.xlsx with Formulas & Styles)
4. Word Document (.docx Styled Letterhead)

Uploads all generated documents to Google Drive via Drive API.
"""

import os
import sys
import json
import shutil
import subprocess
import urllib.parse
import urllib.request
import webbrowser
from pathlib import Path
from datetime import datetime

# UTF-8 stdout configuration for Windows
if hasattr(sys.stdout, 'reconfigure'):
    try:
        sys.stdout.reconfigure(encoding='utf-8', line_buffering=True)
        sys.stderr.reconfigure(encoding='utf-8', line_buffering=True)
    except Exception:
        pass

# Optional document generation imports
try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

try:
    import docx
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import nsdecls, qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload

# Required Google Drive & Gmail scopes
SCOPES = [
    'https://www.googleapis.com/auth/drive.file',
    'https://www.googleapis.com/auth/gmail.send'
]

# -----------------------------------------------------------------------------
# 🏢 MOCK INVOICE DATA (Based on Studio Tunnel Specs & Zomato Ryze Job #91)
# -----------------------------------------------------------------------------
MOCK_INVOICE = {
    "company": {
        "legal_name": "CINELOOM POSTWORKS PRIVATE LIMITED",
        "brand_name": "Studio Tunnel",
        "address": "311, Kamla Spaces, SV Road, Santacruz (West), Mumbai - 400 054",
        "phone": "8928249081",
        "email": "contact@studiotunnel.com",
        "invoice_email": "invoices@studiotunnel.com",
        "gstin": "27AAMCC8604R1ZV",
        "pan": "AAMCC8604R",
        "tan": "PNEC20959B",
        "state": "27-Maharashtra",
        "hsn": "999612"
    },
    "bank": {
        "name": "HDFC Bank",
        "account_no": "50200084321948",
        "ifsc": "HDFC0000079",
        "holder": "CINELOOM POSTWORKS PRIVATE LIMITED",
        "branch": "Santacruz West, Mumbai"
    },
    "client": {
        "name": "ZOMATO MEDIA PRIVATE LIMITED",
        "address": "Ground Floor, 12A, 94 Meghdoot, Nehru Place, New Delhi - 110 019",
        "gstin": "07AAACZ1823G1Z1",
        "pan": "AAACZ1823G",
        "state": "07-Delhi",
        "contact_person": "Production Accounts Team",
        "contact_phone": "+91 98110 00000",
        "contact_email": "accounts@zomato.com"
    },
    "invoice": {
        "number": "INV-91",
        "display_no": "91",
        "date": "01/07/2026",
        "due_date": "16/07/2026",
        "place_of_supply": "07-Delhi",
        "is_inter_state": True,  # MH -> DL = IGST (18%)
        "project_name": "Zomato Ryze Studio Commercial (Director's Cut)",
        "colorist": "Sujith / Yash",
        "line_producer": "Samiran Sonowal"
    },
    "items": [
        {
            "sr_no": 1,
            "description": "Color Grading & Digital Intermediate (DI) Mastering\nProject: Zomato Ryze Commercial (60s + 30s Cutdowns)",
            "hsn": "999612",
            "quantity": 1,
            "unit": "Job",
            "rate": 100000.00,
            "gst_rate": 18,
            "gst_amount": 18000.00,
            "total_amount": 100000.00
        }
    ],
    "financials": {
        "subtotal": 100000.00,
        "cgst": 0.00,
        "sgst": 0.00,
        "igst": 18000.00,
        "total_tax": 18000.00,
        "grand_total": 118000.00,
        "amount_in_words": "One Lakh Eighteen Thousand Rupees Only",
        "received": 0.00,
        "balance": 118000.00
    }
}


# -----------------------------------------------------------------------------
# 1. RENDER HTML INVOICE
# -----------------------------------------------------------------------------
def generate_html_invoice(data, output_path):
    c = data["company"]
    cl = data["client"]
    inv = data["invoice"]
    fin = data["financials"]
    b = data["bank"]

    items_html = ""
    for item in data["items"]:
        desc = item["description"].replace("\n", "<br>")
        items_html += f"""
        <tr>
          <td class="text-center">{item['sr_no']}</td>
          <td><strong>{desc}</strong></td>
          <td class="text-center">{item['hsn']}</td>
          <td class="text-center">{item['quantity']}</td>
          <td class="text-center">{item['unit']}</td>
          <td class="text-right">₹ {item['rate']:,.2f}</td>
          <td class="text-right">₹ {item['gst_amount']:,.2f}<br><small style="color:#64748b;">({item['gst_rate']}%)</small></td>
          <td class="text-right">₹ {item['total_amount']:,.2f}</td>
        </tr>
        """

    tax_rows = ""
    if inv["is_inter_state"]:
        tax_rows = f"""
        <tr>
          <td>IGST @ 18%</td>
          <td class="text-right">₹ {fin['igst']:,.2f}</td>
        </tr>
        """
    else:
        tax_rows = f"""
        <tr>
          <td>CGST @ 9%</td>
          <td class="text-right">₹ {fin['cgst']:,.2f}</td>
        </tr>
        <tr>
          <td>SGST @ 9%</td>
          <td class="text-right">₹ {fin['sgst']:,.2f}</td>
        </tr>
        """

    html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <title>Tax Invoice #{inv['display_no']} - {c['brand_name']}</title>
  <link href="https://fonts.googleapis.com/css2?family=Lexend:wght@300;400;500;600;700&display=swap" rel="stylesheet">
  <style>
    @page {{
      size: A4;
      margin: 12mm 12mm 12mm 12mm;
    }}
    body {{
      font-family: 'Lexend', -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
      color: #1A1A1A;
      font-size: 10.5px;
      line-height: 1.45;
      margin: 0;
      padding: 24px;
      background-color: #FFFFFF;
    }}
    .container {{
      max-width: 800px;
      margin: 0 auto;
    }}
    .header-table {{
      width: 100%;
      margin-bottom: 12px;
      border-collapse: collapse;
    }}
    .company-title {{
      font-size: 15px;
      font-weight: 700;
      color: #1A1A1A;
      letter-spacing: 0.3px;
      margin-bottom: 3px;
    }}
    .brand-badge {{
      background: #1A1A1A;
      color: #00E599;
      display: inline-block;
      padding: 8px 16px;
      font-weight: 700;
      font-size: 15px;
      border-radius: 4px;
      letter-spacing: 1px;
    }}
    .divider-bar {{
      height: 3px;
      background: linear-gradient(90deg, #1A1A1A 0%, #00E599 100%);
      margin: 10px 0 16px 0;
    }}
    .title-banner {{
      font-size: 18px;
      font-weight: 700;
      color: #1A1A1A;
      text-align: center;
      text-transform: uppercase;
      letter-spacing: 1.5px;
      margin-bottom: 16px;
    }}
    .meta-table {{
      width: 100%;
      margin-bottom: 18px;
      border-collapse: collapse;
    }}
    .meta-table td {{
      vertical-align: top;
    }}
    .card-box {{
      background: #F8FAFC;
      border: 1px solid #E2E8F0;
      border-radius: 6px;
      padding: 12px;
    }}
    .section-label {{
      font-weight: 700;
      font-size: 11px;
      color: #0F172A;
      text-transform: uppercase;
      letter-spacing: 0.5px;
      margin-bottom: 6px;
      border-bottom: 2px solid #00E599;
      display: inline-block;
      padding-bottom: 2px;
    }}
    .item-table {{
      width: 100%;
      border-collapse: collapse;
      margin-bottom: 16px;
    }}
    .item-table th {{
      background-color: #1A1A1A;
      color: #CCCCCC;
      font-weight: 600;
      padding: 8px;
      text-align: left;
      font-size: 10px;
    }}
    .item-table td {{
      padding: 8px;
      border-bottom: 1px solid #E2E8F0;
      font-size: 10px;
    }}
    .text-right {{ text-align: right; }}
    .text-center {{ text-align: center; }}
    .summary-wrap {{
      width: 100%;
      margin-top: 10px;
      display: flex;
      justify-content: space-between;
    }}
    .words-box {{
      width: 48%;
      float: left;
      font-size: 10.5px;
    }}
    .totals-table {{
      width: 46%;
      float: right;
      border-collapse: collapse;
    }}
    .totals-table td {{
      padding: 6px 8px;
      font-size: 10.5px;
    }}
    .grand-total-row {{
      background-color: #1A1A1A;
      color: #00E599 !important;
      font-weight: 700;
      font-size: 12px !important;
    }}
    .grand-total-row td {{
      color: #00E599 !important;
    }}
    .clearfix {{ clear: both; }}
    .footer-table {{
      width: 100%;
      margin-top: 28px;
      border-top: 1px solid #E2E8F0;
      padding-top: 16px;
      border-collapse: collapse;
    }}
    .signature-line {{
      border-bottom: 1.5px solid #1A1A1A;
      display: inline-block;
      padding: 0 25px;
      font-weight: 600;
      font-size: 14px;
      margin-top: 20px;
    }}
  </style>
</head>
<body>
<div class="container">

  <!-- Header -->
  <table class="header-table">
    <tr>
      <td style="width: 65%; vertical-align: top;">
        <div class="company-title">{c['legal_name']}</div>
        <div>{c['address']}</div>
        <div><strong>Phone:</strong> {c['phone']} | <strong>Email:</strong> {c['email']}</div>
        <div><strong>GSTIN:</strong> {c['gstin']} | <strong>PAN:</strong> {c['pan']} | <strong>State:</strong> {c['state']}</div>
        <div><strong>TAN:</strong> {c['tan']} | <strong>Default SAC:</strong> {c['hsn']}</div>
      </td>
      <td style="width: 35%; text-align: right; vertical-align: top;">
        <div class="brand-badge">STUDIO TUNNEL</div>
        <div style="font-size: 9.5px; color: #64748b; margin-top: 6px;">Color Grading &bull; DI &bull; Finishing</div>
      </td>
    </tr>
  </table>

  <div class="divider-bar"></div>
  <div class="title-banner">TAX INVOICE</div>

  <!-- Metadata Cards -->
  <table class="meta-table">
    <tr>
      <td style="width: 52%; padding-right: 10px;">
        <div class="card-box">
          <div class="section-label">Bill To</div>
          <div style="font-weight: 700; font-size: 12px; color: #0F172A; margin-bottom: 3px;">{cl['name']}</div>
          <div>{cl['address']}</div>
          <div><strong>GSTIN:</strong> {cl['gstin']}</div>
          <div><strong>PAN:</strong> {cl['pan']} | <strong>State:</strong> {cl['state']}</div>
          <div><strong>Contact:</strong> {cl['contact_person']} ({cl['contact_email']})</div>
        </div>
      </td>
      <td style="width: 48%; padding-left: 10px;">
        <div class="card-box">
          <div class="section-label">Invoice Details</div>
          <div><strong>Invoice No:</strong> <span style="font-size:12px; font-weight:700; color:#0F172A;">{inv['number']}</span></div>
          <div><strong>Invoice Date:</strong> {inv['date']}</div>
          <div><strong>Due Date:</strong> {inv['due_date']} (Net 15 Days)</div>
          <div><strong>Place of Supply:</strong> {inv['place_of_supply']}</div>
          <div><strong>Project:</strong> {inv['project_name']}</div>
          <div><strong>Colorist:</strong> {inv['colorist']}</div>
        </div>
      </td>
    </tr>
  </table>

  <!-- Line Items -->
  <table class="item-table">
    <thead>
      <tr>
        <th style="width: 5%;" class="text-center">#</th>
        <th style="width: 40%;">Description of Services</th>
        <th style="width: 10%;" class="text-center">SAC</th>
        <th style="width: 8%;" class="text-center">Qty</th>
        <th style="width: 7%;" class="text-center">Unit</th>
        <th style="width: 15%;" class="text-right">Rate</th>
        <th style="width: 15%;" class="text-right">GST</th>
        <th style="width: 15%;" class="text-right">Amount</th>
      </tr>
    </thead>
    <tbody>
      {items_html}
    </tbody>
  </table>

  <!-- Summary Section -->
  <div style="width: 100%;">
    <div class="words-box">
      <div style="font-weight: 700; color: #0F172A; margin-bottom: 2px;">Invoice Amount in Words:</div>
      <div style="color: #334155; font-style: italic; margin-bottom: 16px;">{fin['amount_in_words']}</div>

      <div style="font-weight: 700; color: #0F172A; margin-bottom: 2px;">Terms & Notes:</div>
      <div style="color: #64748b; font-size: 9.5px;">
        1. Payment due within 15 days of invoice date.<br>
        2. Electronic remittance preferred. Please quote invoice number during transfer.<br>
        3. Subject to Mumbai Jurisdiction.
      </div>
    </div>

    <table class="totals-table">
      <tr>
        <td>Sub Total</td>
        <td class="text-right">₹ {fin['subtotal']:,.2f}</td>
      </tr>
      {tax_rows}
      <tr class="grand-total-row">
        <td>Grand Total</td>
        <td class="text-right">₹ {fin['grand_total']:,.2f}</td>
      </tr>
      <tr>
        <td>Amount Received</td>
        <td class="text-right">₹ {fin['received']:,.2f}</td>
      </tr>
      <tr style="font-weight:700; border-top:1px solid #CBD5E1;">
        <td>Balance Due</td>
        <td class="text-right" style="color:#0F172A;">₹ {fin['balance']:,.2f}</td>
      </tr>
    </table>
    <div class="clearfix"></div>
  </div>

  <!-- Footer Banking & Signatory -->
  <table class="footer-table">
    <tr>
      <td style="width: 55%; vertical-align: top;">
        <div style="font-weight: 700; color: #0F172A; margin-bottom: 4px;">Electronic Bank Transfer Details:</div>
        <div><strong>Bank Name:</strong> {b['name']}</div>
        <div><strong>Account Number:</strong> <span style="font-family:monospace; font-weight:700; font-size:11px;">{b['account_no']}</span></div>
        <div><strong>IFSC Code:</strong> {b['ifsc']}</div>
        <div><strong>Account Name:</strong> {b['holder']}</div>
        <div><strong>Branch:</strong> {b['branch']}</div>
      </td>
      <td style="width: 45%; text-align: right; vertical-align: top;">
        <div>For <strong>{c['legal_name']}</strong></div>
        <div class="signature-line">{inv['line_producer']}</div><br>
        <div style="color: #64748b; font-size: 9.5px; margin-top: 4px;">Authorized Signatory</div>
      </td>
    </tr>
  </table>

</div>
</body>
</html>
"""
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(html_content)
    print(f"  [OK] Rendered HTML Invoice -> {output_path}")
    return output_path


# -----------------------------------------------------------------------------
# 2. RENDER VECTOR PDF INVOICE (Via Headless Chrome/Edge or ReportLab)
# -----------------------------------------------------------------------------
def generate_pdf_invoice(html_path, output_path):
    browsers = [
        "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome",
        "/Applications/Microsoft Edge.app/Contents/MacOS/Microsoft Edge",
        "/Applications/Brave Browser.app/Contents/MacOS/Brave Browser",
        "/Applications/Chromium.app/Contents/MacOS/Chromium",
        r"C:\Program Files\Google\Chrome\Application\chrome.exe",
        r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe",
        r"C:\Program Files\Microsoft\Edge\Application\msedge.exe",
        r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
        os.path.expandvars(r"%LOCALAPPDATA%\Google\Chrome\Application\chrome.exe"),
        os.path.expandvars(r"%PROGRAMFILES%\Google\Chrome\Application\chrome.exe"),
        shutil.which("google-chrome"),
        shutil.which("chrome"),
        shutil.which("chromium"),
        shutil.which("chromium-browser"),
        shutil.which("msedge")
    ]
    browser_bin = next((b for b in browsers if b and os.path.exists(b)), None)

    if browser_bin:
        abs_html = os.path.abspath(html_path)
        abs_pdf = os.path.abspath(output_path)
        cmd = [
            browser_bin,
            "--headless",
            "--disable-gpu",
            "--no-pdf-header-footer",
            f"--print-to-pdf={abs_pdf}",
            abs_html
        ]
        try:
            subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=15)
            if os.path.exists(output_path):
                print(f"  [OK] Vector PDF Invoice -> {output_path}")
                return output_path
        except Exception as e:
            print(f"  [WARN] Headless browser PDF failed: {e}")

    # Fallback to ReportLab if browser is unavailable
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

        doc = SimpleDocTemplate(
            str(output_path),
            pagesize=A4,
            leftMargin=36,
            rightMargin=36,
            topMargin=36,
            bottomMargin=36
        )
        styles = getSampleStyleSheet()

        c = MOCK_INVOICE["company"]
        cl = MOCK_INVOICE["client"]
        inv = MOCK_INVOICE["invoice"]
        fin = MOCK_INVOICE["financials"]
        b = MOCK_INVOICE["bank"]

        style_normal = ParagraphStyle('NormalText', parent=styles['Normal'], fontSize=9, leading=12, textColor=colors.HexColor("#1A1A1A"))
        style_bold = ParagraphStyle('BoldText', parent=style_normal, fontName='Helvetica-Bold')
        style_title = ParagraphStyle('DocTitle', parent=styles['Title'], fontName='Helvetica-Bold', fontSize=16, leading=20, alignment=1, textColor=colors.HexColor("#1A1A1A"))
        style_small = ParagraphStyle('SmallText', parent=style_normal, fontSize=8, leading=10, textColor=colors.HexColor("#64748B"))
        style_header_dark = ParagraphStyle('HDark', parent=style_bold, fontSize=11, textColor=colors.HexColor("#00E599"))

        story = []

        # Header Table
        header_data = [
            [
                Paragraph(f"<b>{c['legal_name']}</b><br/><font size=8 color='#64748b'>{c['address']}<br/>Phone: {c['phone']} | Email: {c['email']}<br/>GSTIN: {c['gstin']} | PAN: {c['pan']} | State: {c['state']}</font>", style_normal),
                Paragraph("<b>STUDIO TUNNEL</b><br/><font size=7 color='#cccccc'>Color Grading &bull; DI</font>", style_header_dark)
            ]
        ]
        t_header = Table(header_data, colWidths=[380, 140])
        t_header.setStyle(TableStyle([
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
            ('BACKGROUND', (1, 0), (1, 0), colors.HexColor("#1A1A1A")),
            ('TEXTCOLOR', (1, 0), (1, 0), colors.HexColor("#00E599")),
            ('PADDING', (1, 0), (1, 0), 8),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ]))
        story.append(t_header)
        story.append(Spacer(1, 8))
        story.append(HRFlowable(width="100%", thickness=2, color=colors.HexColor("#1A1A1A"), spaceAfter=10))

        story.append(Paragraph("TAX INVOICE", style_title))
        story.append(Spacer(1, 10))

        # Meta Cards (Bill To & Invoice Details)
        meta_data = [
            [
                Paragraph(f"<b>BILL TO:</b><br/><b>{cl['name']}</b><br/>{cl['address']}<br/>GSTIN: {cl['gstin']} | PAN: {cl['pan']}<br/>State: {cl['state']}<br/>Contact: {cl['contact_person']}", style_normal),
                Paragraph(f"<b>INVOICE DETAILS:</b><br/><b>Invoice No: {inv['number']}</b><br/>Date: {inv['date']}<br/>Due Date: {inv['due_date']}<br/>Place of Supply: {inv['place_of_supply']}<br/>Project: {inv['project_name']}", style_normal)
            ]
        ]
        t_meta = Table(meta_data, colWidths=[255, 265])
        t_meta.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor("#F8FAFC")),
            ('BOX', (0, 0), (-1, -1), 0.5, colors.HexColor("#E2E8F0")),
            ('INNERGRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#E2E8F0")),
            ('PADDING', (0, 0), (-1, -1), 8),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ]))
        story.append(t_meta)
        story.append(Spacer(1, 12))

        # Items Table
        items_data = [
            [Paragraph("<b>#</b>", style_small), Paragraph("<b>Description of Services</b>", style_small), Paragraph("<b>SAC</b>", style_small), Paragraph("<b>Qty</b>", style_small), Paragraph("<b>Unit</b>", style_small), Paragraph("<b>Rate (Rs.)</b>", style_small), Paragraph("<b>GST (Rs.)</b>", style_small), Paragraph("<b>Total (Rs.)</b>", style_small)]
        ]
        for item in MOCK_INVOICE["items"]:
            items_data.append([
                Paragraph(str(item["sr_no"]), style_normal),
                Paragraph(item["description"].replace("\n", "<br/>"), style_normal),
                Paragraph(str(item["hsn"]), style_normal),
                Paragraph(str(item["quantity"]), style_normal),
                Paragraph(item["unit"], style_normal),
                Paragraph(f"{item['rate']:,.2f}", style_normal),
                Paragraph(f"{item['gst_amount']:,.2f}", style_normal),
                Paragraph(f"{item['total_amount']:,.2f}", style_normal)
            ])

        t_items = Table(items_data, colWidths=[20, 180, 45, 30, 35, 70, 65, 75])
        t_items.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#1A1A1A")),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor("#CCCCCC")),
            ('ALIGN', (0, 0), (0, -1), 'CENTER'),
            ('ALIGN', (2, 0), (4, -1), 'CENTER'),
            ('ALIGN', (5, 0), (-1, -1), 'RIGHT'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#E2E8F0")),
            ('PADDING', (0, 0), (-1, -1), 6),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ]))
        story.append(t_items)
        story.append(Spacer(1, 10))

        # Totals Table
        totals_data = [
            ["Sub Total", f"Rs. {fin['subtotal']:,.2f}"],
            ["IGST @ 18%", f"Rs. {fin['igst']:,.2f}"],
            ["Grand Total", f"Rs. {fin['grand_total']:,.2f}"],
            ["Amount Received", f"Rs. {fin['received']:,.2f}"],
            ["Balance Due", f"Rs. {fin['balance']:,.2f}"]
        ]
        t_totals = Table(totals_data, colWidths=[120, 100])
        t_totals.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'RIGHT'),
            ('BACKGROUND', (0, 2), (1, 2), colors.HexColor("#1A1A1A")),
            ('TEXTCOLOR', (0, 2), (1, 2), colors.HexColor("#00E599")),
            ('FONTNAME', (0, 2), (1, 2), 'Helvetica-Bold'),
            ('PADDING', (0, 0), (-1, -1), 4),
            ('LINEBELOW', (0, 4), (1, 4), 1, colors.HexColor("#1A1A1A")),
        ]))

        # Words & Bank Table
        words_paragraph = Paragraph(f"<b>Amount in Words:</b><br/><i>{fin['amount_in_words']}</i><br/><br/><b>Payment Details:</b><br/>Bank: {b['name']} | A/C: {b['account_no']}<br/>IFSC: {b['ifsc']} | Branch: {b['branch']}", style_normal)
        t_bottom = Table([[words_paragraph, t_totals]], colWidths=[300, 220])
        t_bottom.setStyle(TableStyle([
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ]))
        story.append(t_bottom)
        story.append(Spacer(1, 20))

        # Signature
        sig_data = [
            ["", f"For {c['legal_name']}\n\n\n_______________________\nSamiran Sonowal\nAuthorized Signatory"]
        ]
        t_sig = Table(sig_data, colWidths=[320, 200])
        t_sig.setStyle(TableStyle([
            ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
            ('VALIGN', (0, 0), (-1, -1), 'BOTTOM'),
        ]))
        story.append(t_sig)

        doc.build(story)
        print(f"  [OK] Vector PDF Invoice (ReportLab Enhanced) -> {output_path}")
        return output_path
    except Exception as e:
        print(f"  [FAIL] Could not generate PDF: {e}")
        return None


# -----------------------------------------------------------------------------
# 3. RENDER EXCEL SPREADSHEET (.xlsx)
# -----------------------------------------------------------------------------
def generate_excel_invoice(data, output_path):
    if not HAS_OPENPYXL:
        print("  [SKIP] openpyxl not installed.")
        return None

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Tax Invoice"

    # Ensure grid lines are visible
    ws.views.sheetView[0].showGridLines = True

    c = data["company"]
    cl = data["client"]
    inv = data["invoice"]
    fin = data["financials"]
    b = data["bank"]

    # Theme Fills and Fonts
    dark_fill = PatternFill(start_color="1A1A1A", end_color="1A1A1A", fill_type="solid")
    card_fill = PatternFill(start_color="F8FAFC", end_color="F8FAFC", fill_type="solid")
    accent_green_font = Font(name="Lexend", size=13, bold=True, color="00E599")
    header_title_font = Font(name="Lexend", size=14, bold=True, color="1A1A1A")
    header_col_font = Font(name="Lexend", size=10, bold=True, color="CCCCCC")
    section_label_font = Font(name="Lexend", size=10, bold=True, color="0F172A")
    bold_dark = Font(name="Lexend", size=10, bold=True, color="0F172A")
    regular_font = Font(name="Lexend", size=10, color="1A1A1A")
    small_gray = Font(name="Lexend", size=9, color="64748B")
    italic_gray = Font(name="Lexend", size=9, italic=True, color="475569")

    thin_border = Border(
        left=Side(style='thin', color='CBD5E1'),
        right=Side(style='thin', color='CBD5E1'),
        top=Side(style='thin', color='CBD5E1'),
        bottom=Side(style='thin', color='CBD5E1')
    )

    # 1. Company Header
    ws.merge_cells("A1:E1")
    ws["A1"] = c["legal_name"]
    ws["A1"].font = header_title_font

    ws.merge_cells("F1:H1")
    ws["F1"] = "STUDIO TUNNEL"
    ws["F1"].fill = dark_fill
    ws["F1"].font = accent_green_font
    ws["F1"].alignment = Alignment(horizontal="center", vertical="center")

    ws["A2"] = c["address"]
    ws["A2"].font = small_gray
    ws["A3"] = f"Phone: {c['phone']} | Email: {c['email']} | GSTIN: {c['gstin']} | State: {c['state']}"
    ws["A3"].font = small_gray
    ws["A4"] = f"PAN: {c['pan']} | TAN: {c['tan']} | Default SAC: {c['hsn']}"
    ws["A4"].font = small_gray

    # 2. Document Title
    ws.merge_cells("A6:H6")
    ws["A6"] = "TAX INVOICE"
    ws["A6"].font = Font(name="Lexend", size=15, bold=True, color="0F172A")
    ws["A6"].alignment = Alignment(horizontal="center")

    # 3. Bill To & Invoice Meta Grid
    ws["A8"] = "BILL TO:"
    ws["A8"].font = section_label_font
    ws["A9"] = cl["name"]
    ws["A9"].font = Font(name="Lexend", size=11, bold=True, color="0F172A")
    ws["A10"] = cl["address"]
    ws["A10"].font = regular_font
    ws["A11"] = f"GSTIN: {cl['gstin']} | PAN: {cl['pan']} | State: {cl['state']}"
    ws["A11"].font = regular_font
    ws["A12"] = f"Contact: {cl['contact_person']} ({cl['contact_email']})"
    ws["A12"].font = regular_font

    ws["F8"] = "INVOICE DETAILS:"
    ws["F8"].font = section_label_font
    ws["F9"] = f"Invoice No: {inv['number']}"
    ws["F9"].font = Font(name="Lexend", size=11, bold=True, color="0F172A")
    ws["F10"] = f"Invoice Date: {inv['date']} | Due Date: {inv['due_date']}"
    ws["F10"].font = regular_font
    ws["F11"] = f"Place of Supply: {inv['place_of_supply']}"
    ws["F11"].font = regular_font
    ws["F12"] = f"Project: {inv['project_name']} (Colorist: {inv['colorist']})"
    ws["F12"].font = regular_font

    # Style Metadata Card Cells with light gray fill
    for r in range(8, 13):
        for c_idx in range(1, 5):
            ws.cell(row=r, column=c_idx).fill = card_fill
        for c_idx in range(6, 9):
            ws.cell(row=r, column=c_idx).fill = card_fill

    # 4. Item Table Header
    headers = ["#", "Description of Services", "SAC", "Qty", "Unit", "Rate (₹)", "GST Amount (₹)", "Total (₹)"]
    header_row = 14
    for col_idx, header in enumerate(headers, 1):
        cell = ws.cell(row=header_row, column=col_idx, value=header)
        cell.fill = dark_fill
        cell.font = header_col_font
        cell.alignment = Alignment(horizontal="center" if col_idx in [1, 3, 4, 5] else ("right" if col_idx >= 6 else "left"), vertical="center")

    # 5. Items Loop
    first_item_row = 15
    current_row = first_item_row
    for item in data["items"]:
        ws.cell(row=current_row, column=1, value=item["sr_no"]).alignment = Alignment(horizontal="center", vertical="top")

        desc_cell = ws.cell(row=current_row, column=2, value=item["description"])
        desc_cell.font = regular_font
        desc_cell.alignment = Alignment(wrap_text=True, vertical="top")

        ws.cell(row=current_row, column=3, value=item["hsn"]).alignment = Alignment(horizontal="center", vertical="top")
        ws.cell(row=current_row, column=4, value=item["quantity"]).alignment = Alignment(horizontal="center", vertical="top")
        ws.cell(row=current_row, column=5, value=item["unit"]).alignment = Alignment(horizontal="center", vertical="top")

        cell_rate = ws.cell(row=current_row, column=6, value=item["rate"])
        cell_rate.number_format = '₹ #,##,##0.00'
        cell_rate.alignment = Alignment(horizontal="right", vertical="top")

        # Dynamic GST Formula
        cell_gst = ws.cell(row=current_row, column=7, value=f"=F{current_row}*({item['gst_rate']}/100)")
        cell_gst.number_format = '₹ #,##,##0.00'
        cell_gst.alignment = Alignment(horizontal="right", vertical="top")

        # Dynamic Total Formula
        cell_tot = ws.cell(row=current_row, column=8, value=f"=F{current_row}+G{current_row}")
        cell_tot.number_format = '₹ #,##,##0.00'
        cell_tot.alignment = Alignment(horizontal="right", vertical="top")

        for c_idx in range(1, 9):
            ws.cell(row=current_row, column=c_idx).border = thin_border
        current_row += 1

    last_item_row = current_row - 1

    # 6. Totals Block
    current_row += 1
    # Subtotal
    ws.cell(row=current_row, column=6, value="Sub Total").font = bold_dark
    ws.cell(row=current_row, column=6).alignment = Alignment(horizontal="right")
    cell_sub = ws.cell(row=current_row, column=8, value=f"=SUM(F{first_item_row}:F{last_item_row})")
    cell_sub.number_format = '₹ #,##,##0.00'
    cell_sub.font = bold_dark
    cell_sub.alignment = Alignment(horizontal="right")
    ws.cell(row=current_row, column=8).border = thin_border

    current_row += 1
    # Tax Row
    tax_label = "IGST @ 18%" if inv["is_inter_state"] else "GST Total (CGST+SGST @ 18%)"
    ws.cell(row=current_row, column=6, value=tax_label).font = bold_dark
    ws.cell(row=current_row, column=6).alignment = Alignment(horizontal="right")
    cell_tax = ws.cell(row=current_row, column=8, value=f"=SUM(G{first_item_row}:G{last_item_row})")
    cell_tax.number_format = '₹ #,##,##0.00'
    cell_tax.font = bold_dark
    cell_tax.alignment = Alignment(horizontal="right")
    ws.cell(row=current_row, column=8).border = thin_border

    current_row += 1
    # Grand Total
    ws.cell(row=current_row, column=6, value="Grand Total").fill = dark_fill
    ws.cell(row=current_row, column=6).font = Font(name="Lexend", size=11, bold=True, color="00E599")
    ws.cell(row=current_row, column=6).alignment = Alignment(horizontal="right")

    cell_gt = ws.cell(row=current_row, column=8, value=f"=SUM(H{first_item_row}:H{last_item_row})")
    cell_gt.fill = dark_fill
    cell_gt.font = Font(name="Lexend", size=11, bold=True, color="00E599")
    cell_gt.number_format = '₹ #,##,##0.00'
    cell_gt.alignment = Alignment(horizontal="right")
    ws.cell(row=current_row, column=8).border = thin_border

    current_row += 1
    # Amount Received
    ws.cell(row=current_row, column=6, value="Amount Received").font = regular_font
    ws.cell(row=current_row, column=6).alignment = Alignment(horizontal="right")
    cell_rec = ws.cell(row=current_row, column=8, value=fin['received'])
    cell_rec.number_format = '₹ #,##,##0.00'
    cell_rec.font = regular_font
    cell_rec.alignment = Alignment(horizontal="right")

    current_row += 1
    # Balance Due
    ws.cell(row=current_row, column=6, value="Balance Due").font = Font(name="Lexend", size=11, bold=True, color="0F172A")
    ws.cell(row=current_row, column=6).alignment = Alignment(horizontal="right")
    cell_bal = ws.cell(row=current_row, column=8, value=f"=H{current_row-2}-H{current_row-1}")
    cell_bal.number_format = '₹ #,##,##0.00'
    cell_bal.font = Font(name="Lexend", size=11, bold=True, color="0F172A")
    cell_bal.alignment = Alignment(horizontal="right")
    ws.cell(row=current_row, column=8).border = thin_border

    # 7. Amount in Words & Notes Section
    current_row += 2
    ws.merge_cells(f"A{current_row}:E{current_row}")
    ws[f"A{current_row}"] = f"Invoice Amount in Words: {fin['amount_in_words']}"
    ws[f"A{current_row}"].font = italic_gray
    ws[f"A{current_row}"].fill = card_fill

    current_row += 2
    ws[f"A{current_row}"] = "BANK TRANSFER DETAILS:"
    ws[f"A{current_row}"].font = section_label_font
    ws[f"F{current_row}"] = f"For {c['legal_name']}"
    ws[f"F{current_row}"].font = bold_dark

    current_row += 1
    ws[f"A{current_row}"] = f"Bank Name: {b['name']} | Account No: {b['account_no']} | IFSC: {b['ifsc']}"
    ws[f"A{current_row}"].font = regular_font
    ws[f"F{current_row}"] = f"{inv['line_producer']} (Authorized Signatory)"
    ws[f"F{current_row}"].font = regular_font

    current_row += 1
    ws[f"A{current_row}"] = f"Account Holder: {b['holder']} | Branch: {b['branch']}"
    ws[f"A{current_row}"].font = regular_font

    # Column Width Optimizations
    ws.column_dimensions["A"].width = 6
    ws.column_dimensions["B"].width = 44
    ws.column_dimensions["C"].width = 12
    ws.column_dimensions["D"].width = 8
    ws.column_dimensions["E"].width = 8
    ws.column_dimensions["F"].width = 18
    ws.column_dimensions["G"].width = 18
    ws.column_dimensions["H"].width = 20

    wb.save(output_path)
    print(f"  [OK] Excel Spreadsheet Invoice -> {output_path}")
    return output_path


# -----------------------------------------------------------------------------
# XML HELPERS FOR DOCX FORMATTING
# -----------------------------------------------------------------------------
def set_docx_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    tcPr.append(shd)

def set_docx_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_docx_cell_borders(cell, top="CBD5E1", bottom="CBD5E1", left=None, right=None, sz="4"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    borders = {'top': top, 'bottom': bottom, 'left': left, 'right': right}
    for b_name, b_val in borders.items():
        if b_val is not None:
            node = OxmlElement(f'w:{b_name}')
            node.set(qn('w:val'), 'single')
            node.set(qn('w:sz'), str(sz))
            node.set(qn('w:space'), '0')
            node.set(qn('w:color'), b_val)
            tcBorders.append(node)
        else:
            node = OxmlElement(f'w:{b_name}')
            node.set(qn('w:val'), 'nil')
            tcBorders.append(node)
    tcPr.append(tcBorders)


# -----------------------------------------------------------------------------
# 4. RENDER WORD DOCUMENT (.docx)
# -----------------------------------------------------------------------------
def generate_docx_invoice(data, output_path):
    if not HAS_DOCX:
        print("  [SKIP] python-docx not installed.")
        return None

    doc = Document()

    # Set page margins to 0.6 inch (~15mm)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    c = data["company"]
    cl = data["client"]
    inv = data["invoice"]
    fin = data["financials"]
    b = data["bank"]

    # 1. Header Table (Company info on Left, Brand Badge on Right)
    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_left, c_right = header_table.rows[0].cells
    c_left.width = Inches(4.8)
    c_right.width = Inches(2.2)

    p_left = c_left.paragraphs[0]
    r_cname = p_left.add_run(f"{c['legal_name']}\n")
    r_cname.bold = True
    r_cname.font.size = Pt(12)
    r_cname.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    r_cinfo = p_left.add_run(f"{c['address']}\nPhone: {c['phone']} | Email: {c['email']}\nGSTIN: {c['gstin']} | PAN: {c['pan']} | State: {c['state']}\nTAN: {c['tan']} | SAC: {c['hsn']}")
    r_cinfo.font.size = Pt(8.5)
    r_cinfo.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    p_right = c_right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_badge = p_right.add_run("  STUDIO TUNNEL  \n")
    r_badge.bold = True
    r_badge.font.size = Pt(13)
    r_badge.font.color.rgb = RGBColor(0x00, 0xE5, 0x99)
    set_docx_cell_shading(c_right, "1A1A1A")
    set_docx_cell_margins(c_right, top=120, bottom=120, left=150, right=150)

    r_sub = p_right.add_run("Color Grading • DI • Finishing")
    r_sub.font.size = Pt(8)
    r_sub.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    # 2. Document Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(12)
    r_title = p_title.add_run("TAX INVOICE")
    r_title.bold = True
    r_title.font.size = Pt(16)
    r_title.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    # 3. Metadata Table (Bill To & Invoice Details)
    meta_table = doc.add_table(rows=1, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    m_left, m_right = meta_table.rows[0].cells
    m_left.width = Inches(3.6)
    m_right.width = Inches(3.4)

    set_docx_cell_shading(m_left, "F8FAFC")
    set_docx_cell_shading(m_right, "F8FAFC")
    set_docx_cell_margins(m_left, top=100, bottom=100, left=120, right=120)
    set_docx_cell_margins(m_right, top=100, bottom=100, left=120, right=120)
    set_docx_cell_borders(m_left, top="E2E8F0", bottom="E2E8F0", left="E2E8F0", right="E2E8F0")
    set_docx_cell_borders(m_right, top="E2E8F0", bottom="E2E8F0", left="E2E8F0", right="E2E8F0")

    p_bto = m_left.paragraphs[0]
    p_bto.add_run("BILL TO:\n").bold = True
    p_bto.runs[0].font.size = Pt(9.5)
    p_bto.runs[0].font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    p_bto.add_run(f"{cl['name']}\n").bold = True
    p_bto.runs[1].font.size = Pt(10)
    p_bto.add_run(f"{cl['address']}\nGSTIN: {cl['gstin']} | PAN: {cl['pan']}\nState: {cl['state']}\nContact: {cl['contact_person']}")
    p_bto.runs[2].font.size = Pt(8.5)

    p_idet = m_right.paragraphs[0]
    p_idet.add_run("INVOICE DETAILS:\n").bold = True
    p_idet.runs[0].font.size = Pt(9.5)
    p_idet.runs[0].font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    p_idet.add_run(f"Invoice No: {inv['number']}\n").bold = True
    p_idet.runs[1].font.size = Pt(10)
    p_idet.add_run(f"Date: {inv['date']} | Due Date: {inv['due_date']}\nPlace of Supply: {inv['place_of_supply']}\nProject: {inv['project_name']}\nColorist: {inv['colorist']}")
    p_idet.runs[2].font.size = Pt(8.5)

    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(8)
    p_space.paragraph_format.space_after = Pt(4)

    # 4. Item Table (8 Columns)
    item_table = doc.add_table(rows=1, cols=8)
    item_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = item_table.rows[0].cells
    hdr_titles = ["#", "Description of Services", "SAC", "Qty", "Unit", "Rate (₹)", "GST Amount (₹)", "Total (₹)"]
    col_widths = [Inches(0.3), Inches(2.7), Inches(0.6), Inches(0.4), Inches(0.4), Inches(0.85), Inches(0.85), Inches(0.9)]

    for i, title in enumerate(hdr_titles):
        hdr_cells[i].width = col_widths[i]
        set_docx_cell_shading(hdr_cells[i], "1A1A1A")
        set_docx_cell_margins(hdr_cells[i], top=80, bottom=80, left=60, right=60)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i >= 5 else (WD_ALIGN_PARAGRAPH.CENTER if i in [0, 2, 3, 4] else WD_ALIGN_PARAGRAPH.LEFT)
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    for item in data["items"]:
        row_cells = item_table.add_row().cells
        vals = [
            str(item["sr_no"]),
            item["description"],
            str(item["hsn"]),
            str(item["quantity"]),
            item["unit"],
            f"₹ {item['rate']:,.2f}",
            f"₹ {item['gst_amount']:,.2f}",
            f"₹ {item['total_amount']:,.2f}"
        ]
        for i, val in enumerate(vals):
            row_cells[i].width = col_widths[i]
            set_docx_cell_margins(row_cells[i], top=80, bottom=80, left=60, right=60)
            set_docx_cell_borders(row_cells[i], top="E2E8F0", bottom="E2E8F0", left=None, right=None)
            p = row_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i >= 5 else (WD_ALIGN_PARAGRAPH.CENTER if i in [0, 2, 3, 4] else WD_ALIGN_PARAGRAPH.LEFT)
            run = p.add_run(val)
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    p_sp2 = doc.add_paragraph()
    p_sp2.paragraph_format.space_before = Pt(8)

    # 5. Financial Totals & Words Block
    tot_table = doc.add_table(rows=1, cols=2)
    tot_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_left, t_right = tot_table.rows[0].cells
    t_left.width = Inches(4.2)
    t_right.width = Inches(2.8)

    p_w = t_left.paragraphs[0]
    p_w.add_run("Amount in Words:\n").bold = True
    p_w.runs[0].font.size = Pt(9)
    r_win = p_w.add_run(f"{fin['amount_in_words']}\n\n")
    r_win.italic = True
    r_win.font.size = Pt(9)
    r_win.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    p_w.add_run("Terms & Conditions:\n").bold = True
    p_w.runs[2].font.size = Pt(8.5)
    r_tms = p_w.add_run("1. Payment due within 15 days of invoice date.\n2. Remittance preferred via Electronic Transfer.\n3. Subject to Mumbai Jurisdiction.")
    r_tms.font.size = Pt(8)
    r_tms.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    p_t = t_right.paragraphs[0]
    p_t.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_t.add_run(f"Sub Total: ₹ {fin['subtotal']:,.2f}\n").font.size = Pt(9)
    p_t.add_run(f"IGST @ 18%: ₹ {fin['igst']:,.2f}\n").font.size = Pt(9)

    r_gt = p_t.add_run(f"Grand Total: ₹ {fin['grand_total']:,.2f}\n")
    r_gt.bold = True
    r_gt.font.size = Pt(11)
    r_gt.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    p_t.add_run(f"Amount Received: ₹ {fin['received']:,.2f}\n").font.size = Pt(8.5)
    r_bal = p_t.add_run(f"Balance Due: ₹ {fin['balance']:,.2f}")
    r_bal.bold = True
    r_bal.font.size = Pt(10)

    # 6. Banking & Signatory Block
    p_sp3 = doc.add_paragraph()
    p_sp3.paragraph_format.space_before = Pt(12)

    bank_table = doc.add_table(rows=1, cols=2)
    bank_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    b_left, b_right = bank_table.rows[0].cells
    b_left.width = Inches(4.2)
    b_right.width = Inches(2.8)

    p_bk = b_left.paragraphs[0]
    p_bk.add_run("Bank Transfer Details:\n").bold = True
    p_bk.runs[0].font.size = Pt(9)
    p_bk.add_run(f"Bank Name: {b['name']}\nAccount No: {b['account_no']}\nIFSC Code: {b['ifsc']}\nAccount Holder: {b['holder']}\nBranch: {b['branch']}").font.size = Pt(8.5)

    p_sig = b_right.paragraphs[0]
    p_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_sig.add_run(f"For {c['legal_name']}\n\n\n\n___________________________\nSamiran Sonowal\nAuthorized Signatory").font.size = Pt(8.5)

    doc.save(output_path)
    print(f"  [OK] Word Document Invoice -> {output_path}")
    return output_path


# -----------------------------------------------------------------------------
# 5. GOOGLE DRIVE UPLOAD MODULE
# -----------------------------------------------------------------------------
def get_drive_service(repo_root):
    token_path = repo_root / "credentials" / "private" / "google_drive_token.json"
    gmail_token_path = repo_root / "credentials" / "private" / "gmail_token.json"
    client_secret_path = repo_root / "credentials" / "private" / "client_secret_972643538415-iotqsas6uh5uanjjgdmal16phvfnsvup.apps.googleusercontent.com.json"

    creds = None
    # Check Drive specific token
    if token_path.exists():
        try:
            creds = Credentials.from_authorized_user_file(str(token_path), SCOPES)
        except Exception:
            creds = None

    # Check unified gmail token if scopes match
    if not creds and gmail_token_path.exists():
        try:
            c = Credentials.from_authorized_user_file(str(gmail_token_path))
            if any('drive' in s for s in c.scopes):
                creds = c
        except Exception:
            pass

    if creds and creds.expired and creds.refresh_token:
        try:
            creds.refresh(Request())
            with open(token_path, 'w', encoding='utf-8') as f:
                f.write(creds.to_json())
        except Exception:
            creds = None

    if not creds or not creds.valid:
        if not client_secret_path.exists():
            print("[WARN] Client secret not found; skipping Google Drive upload.")
            return None

        with open(client_secret_path, 'r', encoding='utf-8') as f:
            secret_data = json.load(f)

        web_config = secret_data.get("web", secret_data.get("installed", {}))
        client_id = web_config.get("client_id")
        client_secret = web_config.get("client_secret")
        redirect_uri = "https://developers.google.com/oauthplayground"

        auth_params = {
            "client_id": client_id,
            "redirect_uri": redirect_uri,
            "response_type": "code",
            "scope": "https://www.googleapis.com/auth/drive.file https://www.googleapis.com/auth/gmail.send",
            "access_type": "offline",
            "prompt": "consent"
        }
        auth_url = "https://accounts.google.com/o/oauth2/auth?" + urllib.parse.urlencode(auth_params)

        print("\n" + "=" * 75)
        print("☁️ GOOGLE DRIVE OAUTH AUTHORIZATION REQUIRED")
        print("=" * 75)
        print("1. Opening Google sign-in page to grant Google Drive upload permissions...")
        print("2. Sign in as: lab@studiotunnel.com")
        print(f"👉 Direct Link: {auth_url}")
        print("=" * 75)

        try:
            webbrowser.open(auth_url, new=2)
        except Exception:
            pass

        auth_code = input("👉 Paste Authorization Code from OAuth Playground (or press Enter to skip upload): ").strip()
        if not auth_code:
            print("[INFO] Skipping Google Drive upload. Files are saved locally.")
            return None

        if "code=" in auth_code:
            parsed = urllib.parse.urlparse(auth_code)
            params = urllib.parse.parse_qs(parsed.query)
            if "code" in params:
                auth_code = params["code"][0]

        token_url = "https://oauth2.googleapis.com/token"
        data = urllib.parse.urlencode({
            "code": auth_code,
            "client_id": client_id,
            "client_secret": client_secret,
            "redirect_uri": redirect_uri,
            "grant_type": "authorization_code"
        }).encode('utf-8')

        req = urllib.request.Request(token_url, data=data, method="POST")
        with urllib.request.urlopen(req) as resp:
            token_resp = json.loads(resp.read().decode('utf-8'))

        creds = Credentials(
            token=token_resp.get("access_token"),
            refresh_token=token_resp.get("refresh_token"),
            token_uri=token_url,
            client_id=client_id,
            client_secret=client_secret,
            scopes=SCOPES
        )

        with open(token_path, 'w', encoding='utf-8') as f:
            f.write(creds.to_json())
        print("[OK] Google Drive token saved to credentials/private/google_drive_token.json")

    return build('drive', 'v3', credentials=creds)


def upload_file_to_drive(drive_service, file_path, mime_type):
    file_name = os.path.basename(file_path)
    file_metadata = {'name': file_name}
    media = MediaFileUpload(file_path, mimetype=mime_type, resumable=True)

    print(f"  [DRIVE] Uploading {file_name} ...", flush=True)
    uploaded_file = drive_service.files().create(
        body=file_metadata,
        media_body=media,
        fields='id, webViewLink, webContentLink'
    ).execute()

    return uploaded_file


# -----------------------------------------------------------------------------
# MAIN PIPELINE EXECUTION
# -----------------------------------------------------------------------------
def main():
    repo_root = Path(__file__).resolve().parent.parent
    output_dir = repo_root / "output" / "mock-invoices"
    output_dir.mkdir(parents=True, exist_ok=True)

    base_name = f"INV91_ZOMATO_STUDIOTUNNEL_{datetime.now().strftime('%Y%m%d')}"
    html_file = output_dir / f"{base_name}.html"
    pdf_file = output_dir / f"{base_name}.pdf"
    excel_file = output_dir / f"{base_name}.xlsx"
    docx_file = output_dir / f"{base_name}.docx"

    print("=" * 75)
    print("🚀 ST-IN-gen — MULTI-FORMAT MOCK INVOICE PIPELINE (WIP)")
    print("=" * 75)
    print(f"Invoice Target : {MOCK_INVOICE['invoice']['number']} (Client: {MOCK_INVOICE['client']['name']})")
    print(f"Output Folder  : {output_dir}\n")

    print("📄 1. Generating Multi-Format Documents:")
    generate_html_invoice(MOCK_INVOICE, str(html_file))
    generate_pdf_invoice(str(html_file), str(pdf_file))
    generate_excel_invoice(MOCK_INVOICE, str(excel_file))
    generate_docx_invoice(MOCK_INVOICE, str(docx_file))

    print("\n" + "=" * 75)
    print("☁️ 2. Google Drive Ingestion:")
    print("=" * 75)

    drive_service = get_drive_service(repo_root)
    if drive_service:
        files_to_upload = [
            (html_file, "text/html"),
            (pdf_file, "application/pdf"),
            (excel_file, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"),
            (docx_file, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        ]

        uploaded_links = []
        for f_path, m_type in files_to_upload:
            if os.path.exists(f_path):
                res = upload_file_to_drive(drive_service, str(f_path), m_type)
                uploaded_links.append((f_path.name, res.get('id'), res.get('webViewLink')))

        print("\n" + "=" * 75)
        print("✨ MULTI-FORMAT INVOICES UPLOADED TO GOOGLE DRIVE:")
        print("=" * 75)
        for name, file_id, link in uploaded_links:
            print(f"  • {name:<35} : {link}")
        print("=" * 75)

    print("\n✨ LOCAL FILES READY IN output/mock-invoices/ :")
    print(f"  1. HTML   : {html_file}")
    print(f"  2. PDF    : {pdf_file}")
    print(f"  3. Excel  : {excel_file}")
    print(f"  4. Word   : {docx_file}")
    print("=" * 75)


if __name__ == "__main__":
    sys.exit(main())
